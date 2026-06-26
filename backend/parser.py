import io
from typing import Optional

import fitz
from docx import Document


ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

MIME_MAP = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
}


def extract_text_from_pdf(file_bytes: bytes) -> str:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text_parts: list[str] = []
    for page in doc:
        text_parts.append(page.get_text())
    doc.close()
    return "\n".join(text_parts).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs).strip()


def extract_text_from_txt(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return file_bytes.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace").strip()


def detect_file_type(filename: Optional[str], content_type: Optional[str]) -> str:
    if content_type and content_type in MIME_MAP:
        return MIME_MAP[content_type]

    if filename:
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if ext == "pdf":
            return "pdf"
        if ext == "docx":
            return "docx"
        if ext == "txt":
            return "txt"

    raise ValueError("Unsupported file type. Allowed: PDF, DOCX, TXT.")


def parse_document(file_bytes: bytes, filename: Optional[str], content_type: Optional[str]) -> str:
    if len(file_bytes) > MAX_FILE_SIZE:
        raise ValueError("File exceeds maximum size of 10MB.")

    if len(file_bytes) == 0:
        raise ValueError("Document is empty.")

    file_type = detect_file_type(filename, content_type)

    if file_type == "pdf":
        text = extract_text_from_pdf(file_bytes)
        if not text or len(text.strip()) < 10:
            raise ValueError(
                "No extractable text found. This appears to be a scanned PDF. "
                "Please upload a text-based PDF or use OCR preprocessing."
            )
        return text

    if file_type == "docx":
        text = extract_text_from_docx(file_bytes)
        if not text:
            raise ValueError("Document is empty or contains no readable text.")
        return text

    text = extract_text_from_txt(file_bytes)
    if not text:
        raise ValueError("Document is empty.")
    return text
